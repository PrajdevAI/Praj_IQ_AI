"""File parsing and text extraction utilities for multiple file types."""

import io
import os
import logging
from typing import Dict

import PyPDF2
import pdfplumber

logger = logging.getLogger(__name__)

# =========================================================================
# Supported file extensions and their MIME categories
# =========================================================================
SUPPORTED_EXTENSIONS = {
    # PDF
    ".pdf": "pdf",
    # Word
    ".docx": "docx",
    ".doc": "doc",
    # Excel
    ".xlsx": "xlsx",
    ".xls": "xls",
    # CSV / Text
    ".csv": "csv",
    ".tsv": "csv",
    ".txt": "text",
    ".md": "text",
    ".rst": "text",
    ".log": "text",
    ".json": "text",
    ".xml": "text",
    ".html": "text",
    ".htm": "text",
    ".yaml": "text",
    ".yml": "text",
    # Images (OCR)
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".tiff": "image",
    ".tif": "image",
    ".bmp": "image",
    ".webp": "image",
    ".svg": "svg",
}

# Extensions shown in the file uploader widget
UPLOAD_EXTENSIONS = list(SUPPORTED_EXTENSIONS.keys())


def get_file_type(filename: str) -> str:
    """Get the file type category from filename extension."""
    ext = os.path.splitext(filename.lower())[1]
    return SUPPORTED_EXTENSIONS.get(ext, "unknown")


def validate_file(file_bytes: bytes, filename: str) -> bool:
    """
    Validate if file is a supported and readable format.

    Args:
        file_bytes: File bytes to validate
        filename: Original filename (used to detect type)

    Returns:
        True if valid, False otherwise
    """
    if not file_bytes or len(file_bytes) == 0:
        logger.warning("Empty file")
        return False

    file_type = get_file_type(filename)

    if file_type == "unknown":
        logger.warning(f"Unsupported file type: {filename}")
        return False

    # PDF-specific validation
    if file_type == "pdf":
        return validate_pdf(file_bytes)

    return True


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from any supported file type.

    Args:
        file_bytes: File as bytes
        filename: Original filename (used to detect type)

    Returns:
        Extracted text as string

    Raises:
        ValueError: If file type is unsupported or extraction fails
    """
    file_type = get_file_type(filename)
    logger.info(f"Extracting text from '{filename}' (type: {file_type})")

    try:
        if file_type == "pdf":
            return extract_text_from_pdf(file_bytes)
        elif file_type == "docx":
            return _extract_from_docx(file_bytes)
        elif file_type == "doc":
            return _extract_from_doc(file_bytes)
        elif file_type == "xlsx" or file_type == "xls":
            return _extract_from_excel(file_bytes, filename)
        elif file_type == "csv":
            return _extract_from_csv(file_bytes)
        elif file_type == "text":
            return _extract_from_text(file_bytes)
        elif file_type == "image":
            return _extract_from_image(file_bytes, filename)
        elif file_type == "svg":
            return _extract_from_svg(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Text extraction failed for {filename}: {e}")
        raise ValueError(f"Failed to extract text from {filename}: {str(e)}")


# =========================================================================
# PDF extraction (existing logic preserved)
# =========================================================================

def extract_text_from_pdf(file_bytes: bytes, method: str = "pdfplumber") -> str:
    """Extract text from PDF file bytes."""
    try:
        if method == "pdfplumber":
            return _extract_pdf_with_pdfplumber(file_bytes)
        else:
            return _extract_pdf_with_pypdf2(file_bytes)
    except Exception as e:
        logger.error(f"PDF extraction failed with {method}: {str(e)}")
        if method == "pdfplumber":
            logger.info("Falling back to PyPDF2")
            return _extract_pdf_with_pypdf2(file_bytes)
        else:
            logger.info("Falling back to pdfplumber")
            return _extract_pdf_with_pdfplumber(file_bytes)


def _extract_pdf_with_pdfplumber(file_bytes: bytes) -> str:
    """Extract text using pdfplumber (better for complex layouts)."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num}]\n{text}")
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num}: {str(e)}")
                continue
    return "\n\n".join(text_parts)


def _extract_pdf_with_pypdf2(file_bytes: bytes) -> str:
    """Extract text using PyPDF2 (faster, simpler)."""
    text_parts = []
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page_num, page in enumerate(pdf_reader.pages, 1):
        try:
            text = page.extract_text()
            if text:
                text_parts.append(f"[Page {page_num}]\n{text}")
        except Exception as e:
            logger.warning(f"Failed to extract page {page_num}: {str(e)}")
            continue
    return "\n\n".join(text_parts)


def validate_pdf(file_bytes: bytes) -> bool:
    """Validate if file is a proper PDF."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        _ = len(pdf_reader.pages)
        return True
    except Exception as e:
        logger.warning(f"PDF validation failed: {str(e)}")
        return False


def extract_metadata(file_bytes: bytes) -> Dict:
    """Extract PDF metadata."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        metadata = pdf_reader.metadata or {}
        return {
            "title": metadata.get("/Title", ""),
            "author": metadata.get("/Author", ""),
            "subject": metadata.get("/Subject", ""),
            "creator": metadata.get("/Creator", ""),
            "producer": metadata.get("/Producer", ""),
            "creation_date": metadata.get("/CreationDate", ""),
            "num_pages": len(pdf_reader.pages),
        }
    except Exception as e:
        logger.error(f"Failed to extract metadata: {str(e)}")
        return {}


# =========================================================================
# Word documents (.docx)
# =========================================================================

def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from .docx files using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "python-docx is required for .docx files. "
            "Install it with: pip install python-docx"
        )

    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract text from tables
    for table_idx, table in enumerate(doc.tables, 1):
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            text_parts.append(f"\n[Table {table_idx}]\n" + "\n".join(table_rows))

    result = "\n\n".join(text_parts)
    logger.info(f"Extracted {len(result)} chars from DOCX ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
    return result


def _extract_from_doc(file_bytes: bytes) -> str:
    """
    Extract text from legacy .doc files.
    Tries textract first, falls back to antiword-style extraction.
    """
    # Try using python-docx (works for some .doc files saved as docx internally)
    try:
        return _extract_from_docx(file_bytes)
    except Exception:
        pass

    # Try raw text extraction as last resort
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        # Filter out binary garbage — keep only printable content
        cleaned = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
        if len(cleaned.strip()) > 50:
            logger.info("Extracted text from .doc via raw decode")
            return cleaned
    except Exception:
        pass

    raise ValueError(
        "Cannot extract text from .doc file. "
        "Please convert it to .docx format and re-upload."
    )


# =========================================================================
# Excel files (.xlsx, .xls)
# =========================================================================

def _extract_from_excel(file_bytes: bytes, filename: str) -> str:
    """Extract text from Excel files using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise ValueError(
            "openpyxl is required for Excel files. "
            "Install it with: pip install openpyxl"
        )

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows = []

        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            row_text = " | ".join(cells).strip()
            if row_text and row_text != "| " * len(cells):
                sheet_rows.append(row_text)

        if sheet_rows:
            text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_rows))

    wb.close()
    result = "\n\n".join(text_parts)
    logger.info(f"Extracted {len(result)} chars from Excel ({len(wb.sheetnames)} sheets)")
    return result


# =========================================================================
# CSV / TSV files
# =========================================================================

def _extract_from_csv(file_bytes: bytes) -> str:
    """Extract text from CSV/TSV files."""
    import csv

    # Detect encoding
    text = None
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        text = file_bytes.decode("utf-8", errors="replace")

    # Parse CSV
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        row_text = " | ".join(cell.strip() for cell in row)
        if row_text.strip():
            rows.append(row_text)

    result = "\n".join(rows)
    logger.info(f"Extracted {len(result)} chars from CSV ({len(rows)} rows)")
    return result


# =========================================================================
# Plain text files (.txt, .md, .json, .xml, .html, etc.)
# =========================================================================

def _extract_from_text(file_bytes: bytes) -> str:
    """Extract text from plain text files."""
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            text = file_bytes.decode(encoding)
            logger.info(f"Decoded text file ({encoding}): {len(text)} chars")
            return text
        except UnicodeDecodeError:
            continue

    return file_bytes.decode("utf-8", errors="replace")


# =========================================================================
# Image files (OCR using pytesseract if available, else boto3 Textract)
# =========================================================================

def _extract_from_image(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from images.
    Tries pytesseract (local OCR) first, then falls back to
    AWS Textract if available, then to a simple error message.
    """
    # Strategy 1: pytesseract (local OCR)
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image)
        if text.strip():
            logger.info(f"Extracted {len(text)} chars from image via pytesseract")
            return text
    except ImportError:
        logger.debug("pytesseract not available, trying alternatives")
    except Exception as e:
        logger.warning(f"pytesseract OCR failed: {e}")

    # Strategy 2: AWS Textract
    try:
        import boto3
        from config.settings import settings

        # textract = boto3.client(
        #     "textract",
        #     region_name=settings.AWS_REGION,
        #     aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
        #     aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
        # )
        textract = boto3.client(
            "textract",
            region_name=settings.AWS_REGION,
            )

        response = textract.detect_document_text(Document={"Bytes": file_bytes})
        blocks = response.get("Blocks", [])
        lines = [b["Text"] for b in blocks if b["BlockType"] == "LINE"]

        if lines:
            text = "\n".join(lines)
            logger.info(f"Extracted {len(text)} chars from image via AWS Textract")
            return text
    except ImportError:
        logger.debug("boto3/textract not available")
    except Exception as e:
        logger.warning(f"AWS Textract failed: {e}")

    raise ValueError(
        f"Cannot extract text from image '{filename}'. "
        "Install pytesseract (pip install pytesseract Pillow) for local OCR, "
        "or ensure AWS Textract is configured."
    )


# =========================================================================
# SVG files
# =========================================================================

def _extract_from_svg(file_bytes: bytes) -> str:
    """Extract text content from SVG files."""
    import re

    text = file_bytes.decode("utf-8", errors="replace")

    # Extract text from <text> elements
    text_elements = re.findall(r"<text[^>]*>(.*?)</text>", text, re.DOTALL)
    # Also extract from <tspan> elements
    tspan_elements = re.findall(r"<tspan[^>]*>(.*?)</tspan>", text, re.DOTALL)

    all_text = text_elements + tspan_elements

    if all_text:
        # Clean HTML tags from extracted text
        cleaned = [re.sub(r"<[^>]+>", "", t).strip() for t in all_text]
        result = "\n".join(t for t in cleaned if t)
        logger.info(f"Extracted {len(result)} chars from SVG")
        return result

    # If no text elements, return the raw SVG as text (it may contain useful content)
    logger.warning("No <text> elements found in SVG, returning raw content")
    return text
